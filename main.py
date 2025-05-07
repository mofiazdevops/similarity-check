from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
from sklearn.feature_extraction.text import CountVectorizer
import os
import requests
import base64

# Constants
ASSIGNMENT_DIR = "downloaded_docs/"
REPORTS_DIR = "similarity_reports/"
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}
download_dir = './downloaded_docs'

# Ensure directories exist
os.makedirs(ASSIGNMENT_DIR, exist_ok=True)

os.makedirs(REPORTS_DIR, exist_ok=True)

os.makedirs(download_dir, exist_ok=True)


# Flask app initialization
app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB limit

def download_files(download_dir):
    # Ensure the download directory exists
    os.makedirs(download_dir, exist_ok=True)

    url = "https://staging.portalteam.org/api/files"
    #url = "http://localhost/PortalCRM/api/files"
    response = requests.get(url)
    if response.status_code != 200:
        print("Failed to retrieve the file list")
        return

    files = response.json().get('data', [])
    
    # Download only new files
    for file_info in files:
        file_name = file_info['file_name']
        file_link = file_info['file_link']
        file_path = os.path.join(download_dir, file_name)

        # Check if the file already exists
        if os.path.exists(file_path):
            print(f"Skipping {file_name} (already exists)")
            continue

        # Download and save the new file
        file_response = requests.get(file_link)
        if file_response.status_code == 200:
            with open(file_path, 'wb') as file:
                file.write(file_response.content)
            print(f"Downloaded {file_name}")
        else:
            print(f"Failed to download {file_name}")


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf_old(file_path):
    try:
        reader = PdfReader(file_path)
        return ''.join([page.extract_text() for page in reader.pages])
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

def extract_text_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            try:
                text += page.extract_text() or ''  # Handle pages with no text
            except Exception as e:
                print(f"Skipped a problematic page due to: {str(e).encode('utf-8', 'ignore').decode('utf-8')}")
                continue  # Skip the problematic page
        return text
    except Exception as e:
        print(f"Error processing PDF file: {str(e).encode('utf-8', 'ignore').decode('utf-8')}")
        return None

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        text = ''
        for para in doc.paragraphs:
            try:
                text += para.text + '\n'
            except Exception as e:
                print(f"Skipped a problematic paragraph due to: {e}")
                continue  # Skip the problematic paragraph
        return text
    except Exception as e:
        print(f"Error processing DOCX file: {str(e).encode('utf-8', 'ignore').decode('utf-8')}")
        return None

def extract_text_from_docx_old(file_path):
    try:
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None

# def find_matching_phrases(source_text, target_text, n=5):
#     vectorizer = CountVectorizer(ngram_range=(n, n)).fit([source_text, target_text])
#     source_ngrams = set(vectorizer.build_analyzer()(source_text))
#     target_ngrams = set(vectorizer.build_analyzer()(target_text))
#     matches = list(source_ngrams.intersection(target_ngrams))
#     similarity_percentage = (len(matches) / len(source_ngrams)) * 100 if source_ngrams else 0
#     return matches, similarity_percentage


def find_matching_phrases(source_text, target_text, n=5):
    """
    Find matching phrases of n words between source and target texts.
    Ensure matches are non-overlapping with at least n-word separation.
    """
    # Split text into tokens (words)
    source_words = source_text.split()
    target_words = target_text.split()
    
    # Generate n-grams for source and target texts
    source_ngrams = [
        " ".join(source_words[i:i + n]) 
        for i in range(0, len(source_words) - n + 1)
    ]
    target_ngrams = [
        " ".join(target_words[i:i + n]) 
        for i in range(0, len(target_words) - n + 1)
    ]
    
    matches = []
    used_indices = set()  # To track used indices in the source text for non-overlapping
    
    # Check for matches
    for i, phrase in enumerate(source_ngrams):
        if i in used_indices:  # Skip already matched phrases
            continue
        
        if phrase in target_ngrams:
            matches.append(phrase)
            
            # Mark indices for the current match to avoid duplication
            used_indices.update(range(i, i + n))
    
    # Calculate similarity percentage
    unique_source_phrases = len(source_ngrams)
    similarity_percentage = (len(matches) / unique_source_phrases) * 100 if unique_source_phrases > 0 else 0
    
    return matches, similarity_percentage



def generate_pdf_report_old(source_text, matches, overall_similarity, filename):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Add header
    pdf.set_font("Arial", 'B', size=14)
    pdf.cell(0, 10, "Detailed Similarity Report", ln=True, align="C")
    pdf.ln(10)

    # Add overall similarity
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, f"Overall Similarity: {overall_similarity:.2f}%", ln=True)
    pdf.ln(10)

    # Add source document content
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, "Source Document Content:", border='B')
    pdf.ln(5)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 10, source_text)
    pdf.ln(10)

    # Add matching details
    for match in matches:
        pdf.set_font("Arial", 'B', size=12)
        pdf.multi_cell(0, 10, f"Matching Document: ", border=0, align="L")
        pdf.ln(5)
        pdf.multi_cell(0, 10, f"{match['document_name']} ({match['similarity_percentage']:.2f}% similarity)", border=0, align="L")
        pdf.cell(0, 10, f"View File", border=0, ln=1, align="L", link=match['document'])        
        pdf.ln(5)
        pdf.set_font("Arial", size=10)
        for phrase in match["matching_phrases"]:
            pdf.multi_cell(0, 10, "Matching Phrase: " + phrase, border=0, align="L")
        pdf.ln(5)

    # Save the report
    
    pdf_path = os.path.join(REPORTS_DIR, filename)
    pdf.output(pdf_path)
    return pdf_path

def generate_pdf_report(source_text, matches, overall_similarity, filename):
    from unidecode import unidecode

    def safe_text(text):
        """Converts Unicode characters to ASCII-safe versions."""
        return unidecode(text) if text else text

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Add header
    pdf.set_font("Arial", 'B', size=14)
    pdf.cell(0, 10, safe_text("Detailed Similarity Report"), ln=True, align="C")
    pdf.ln(10)

    # Add overall similarity
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, safe_text(f"Overall Similarity: {overall_similarity:.2f}%"), ln=True)
    pdf.ln(10)

    # Add source document content
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, safe_text("Source Document Content:"), border='B')
    pdf.ln(5)
    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 10, safe_text(source_text))
    pdf.ln(10)
    matches.sort(key=lambda x: x["similarity_percentage"], reverse=True)
    # Add matching details
    for match in matches:
        pdf.set_font("Arial", 'B', size=12)
        pdf.multi_cell(0, 10, safe_text("Matching Document: "), border=0, align="L")
        pdf.ln(5)
        pdf.multi_cell(0, 10, safe_text(f"{match['document_name']} ({match['similarity_percentage']:.2f}% similarity)"), border=0, align="L")
        pdf.cell(0, 10, safe_text("View File"), border=0, ln=1, align="L", link=match['document'])        
        pdf.ln(5)
        pdf.set_font("Arial", size=10)
        for phrase in match["matching_phrases"]:
            pdf.multi_cell(0, 10, safe_text("Matching Phrase: " + phrase), border=0, align="L")
        pdf.ln(5)

    # Save the report
    pdf_path = os.path.join(REPORTS_DIR, filename)
    pdf.output(pdf_path)
    return pdf_path

@app.route('/check-similarity', methods=['POST'])
def check_similarity():
    print("Entered...")
    print(request)
    # Check if request contains JSON
    if not request.is_json:
        print("Request must be JSON")
        return jsonify({"error": "Request must be JSON"}), 400
    try:
        data = request.get_json()
    except Exception as e:
        print("Error:", str(e))
        return jsonify({"error": str(e)}), 500
    #print("Data...", data)
    if not data or 'file' not in data or 'insert_id' not in data:
        print('Missing file or insert_id')
        return jsonify({'error': 'Missing file or insert_id'}), 400
    
    # Decode file data
    file_name = data['file_name']
    file_data = base64.b64decode(data['file'])
    insert_id = data['insert_id']
    BASE_URL = "https://staging.portalteam.org/user_uploads"
    #BASE_URL = "http://localhost/PortalCRM/user_uploads"
    

    if not allowed_file(file_name):
        print('Unsupported file format. Upload a PDF or DOCX file.')
        return jsonify({'error': 'Unsupported file format. Upload a PDF or DOCX file.'}), 400

    
    filename = secure_filename(file_name)
    file_path = os.path.join('./', filename)
    #file.save(file_path)
    try:
        with open(file_path, 'wb') as f:
            f.write(file_data)  # Write the decoded file content
        print(f"File saved successfully: {file_path}")
    except Exception as e:
        print(f"Error saving file: {e}")
        return jsonify({'error': 'Error saving the file.'}), 500

    # Extract source text
    if filename.endswith('.pdf'):
        source_text = extract_text_from_pdf(file_path)
    elif filename.endswith('.docx'):
        source_text = extract_text_from_docx(file_path)
    else:
        print('Unsupported file type')
        return jsonify({'error': 'Unsupported file type.'}), 400

    if not source_text:
        print('Failed to extract text from the uploaded file.')
        return jsonify({'error': 'Failed to extract text from the uploaded file.'}), 400

    # Find matches
    matches = []
    total_similarity = 0
    total_sources = 0

    for existing_file in os.listdir(ASSIGNMENT_DIR):
        if existing_file == filename or not allowed_file(existing_file):
            continue

        existing_file_path = os.path.join(ASSIGNMENT_DIR, existing_file)
        if existing_file.endswith('.pdf'):
            target_text = extract_text_from_pdf(existing_file_path)
        elif existing_file.endswith('.docx'):
            target_text = extract_text_from_docx(existing_file_path)
        else:
            continue

        if target_text:
            if source_text.strip() == target_text.strip():
                matches.append({
                    "document": f"{BASE_URL}/{existing_file}",
                    "document_name" : existing_file,
                    "matching_phrases": ["Exact match with uploaded file"],
                    "similarity_percentage": 100
                })
                total_similarity += 100  # Add 100% similarity for exact file
                total_sources += 1
            else:
                matching_phrases, similarity_percentage = find_matching_phrases(source_text, target_text, n=5)
                if matching_phrases and similarity_percentage > 2:
                    matches.append({
                        "document": f"{BASE_URL}/{existing_file}",
                        "document_name" : existing_file,
                        "matching_phrases": matching_phrases,
                        "similarity_percentage": similarity_percentage
                    })
                    total_similarity += similarity_percentage * len(matching_phrases)
                    total_sources += 1 #len(matching_phrases)

    #overall_similarity = (total_similarity / total_sources) if total_sources else 0
    overall_similarity = (total_similarity / total_sources) if total_sources else 0
    overall_similarity = min(overall_similarity * 100, 100)
    #overall_similarity = min((total_similarity / total_sources) if total_sources else 0, 100)
    print(total_similarity)
    print(total_sources)
    os.remove(file_path)
    print('Source file reomved successfully')
    # if not matches:
    #     matches.append({
    #                 "document": f"No significant matches found.",
    #                 "document_name" : "No significant matches found.",
    #                 "matching_phrases": ["No significant matches found."],
    #                 "similarity_percentage": 0
    #             })

    # Generate report
    report_filename = f"{filename.rsplit('.', 1)[0]}_similarity_report.pdf"

    pdf_path = generate_pdf_report(source_text, matches, overall_similarity, report_filename)
    with open(pdf_path, 'rb') as file:
        file_data = base64.b64encode(file.read()).decode('utf-8')
    callback_url = "https://staging.portalteam.org/api/files/save-response"
    #callback_url = "http://localhost/PortalCRM/api/files/save-response"
    data = {
        'insert_id': insert_id,
        'similarity': overall_similarity,
        'file_name': report_filename,
        'file_data': file_data,
    }
    
    response = requests.post(callback_url, json=data)
    
    download_files(download_dir)
    #print(f"Status Code: {response.status_code}")
    #print(f"Response Headers: {response.headers}")
    #print(f"Response Body: {response.text}")

    # If you want to process the JSON response (if applicable)
    #if response.headers.get('Content-Type') == 'application/json':
        #response_json = response.json()
        #print(f"Response JSON: {response_json}")
    return jsonify({
        'message': 'Similarity report generated successfully.',
        "Total similarity" : f"{total_similarity}",
        "Total Source" : f"{total_sources}",
        'overall_similarity': f"{overall_similarity:.2f}",
        'report_path': pdf_path,
        'matches': [
            {
                'document': match["document"],
                'similarity_percentage': f"{match['similarity_percentage']:.2f}",
            } for match in matches
        ]
    }), 200

if __name__ == '__main__':
    download_files(download_dir)
    app.run(debug=True, host='0.0.0.0', port=8002)
 